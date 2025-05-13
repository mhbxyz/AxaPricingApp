import os

from flask import Blueprint, request, jsonify, send_file, current_app
from datetime import datetime
from docx import Document
from weasyprint import HTML

from axa_pricing_app import db
from axa_pricing_app.models import Quote

bp = Blueprint('quotes', __name__)


@bp.route('/quotes', methods=['GET'])
def get_quotes():
    quotes = Quote.query.all()
    return jsonify([q.serialize() for q in quotes])

@bp.route("/quotes", methods=["POST"])
def create_quote():
    data = request.json

    # Règle métier : Habitation => DO seule uniquement
    if data["destination"].lower() == "habitation" and data["warranty_type"] != "DO seule":
        return jsonify({"error": "Pour une destination 'Habitation', seule la garantie 'DO seule' est autorisée."}), 400

    try:
        quote = Quote(
            opportunity_number=data["opportunity_number"],
            client_name=data["client_name"],
            destination=data["destination"],
            warranty_type=data["warranty_type"],
            warranty_rate=float(data["warranty_rate"]),
            work_type=data["work_type"],
            cost=float(data["cost"])
        )
        db.session.add(quote)
        db.session.commit()

        timestamp = quote.created_at.strftime('%Y-%m-%d_%H-%M')
        filename_base = f"Proposition_commerciale_{quote.opportunity_number}_{timestamp}"
        os.makedirs(current_app.config['UPLOAD_FOLDER'], exist_ok=True)
        docx_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{filename_base}.docx")
        pdf_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{filename_base}.pdf")

        # Calcul de la prime seule
        prime = quote.cost * quote.warranty_rate
        rate_percent = quote.warranty_rate * 100

        # DOCX
        doc = Document()
        doc.add_heading("Proposition Commerciale", level=0)

        doc.add_paragraph(f"Date : {quote.created_at.strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Numéro d'opportunité : {quote.opportunity_number}")
        doc.add_paragraph(f"Client : {quote.client_name}")
        doc.add_paragraph(f"Destination de l’ouvrage : {quote.destination}")
        doc.add_paragraph(f"Type de garantie : {quote.warranty_type}")
        doc.add_paragraph(f"Taux appliqué : {rate_percent:.2f} %")
        doc.add_paragraph(f"Type de travaux : {quote.work_type}")
        doc.add_paragraph(f"Coût de l’ouvrage : {quote.cost:,.2f} €")
        doc.add_paragraph(f"Prime seule : {prime:,.2f} €")

        doc.save(docx_path)

        # PDF
        html_content = f"""
        <h1>Proposition Commerciale</h1>
        <p><strong>Date :</strong> {quote.created_at.strftime('%d/%m/%Y %H:%M')}</p>
        <p><strong>Numéro d'opportunité :</strong> {quote.opportunity_number}</p>
        <p><strong>Client :</strong> {quote.client_name}</p>
        <p><strong>Destination de l’ouvrage :</strong> {quote.destination}</p>
        <p><strong>Type de garantie :</strong> {quote.warranty_type}</p>
        <p><strong>Taux appliqué :</strong> {rate_percent:.2f} %</p>
        <p><strong>Type de travaux :</strong> {quote.work_type}</p>
        <p><strong>Coût de l’ouvrage :</strong> {quote.cost:,.2f} €</p>
        <p><strong>Prime seule :</strong> {prime:,.2f} €</p>
        """

        HTML(string=html_content).write_pdf(pdf_path)

        return jsonify({
            "message": "Quote created",
            "id": quote.id,
            "docx_url": f"/quotes/{quote.id}/document?type=docx",
            "pdf_url": f"/quotes/{quote.id}/document?type=pdf"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@bp.route('/quotes/<int:quote_id>/document')
def download_document(quote_id):
    doc_type = request.args.get('type', 'pdf')
    quote = Quote.query.get_or_404(quote_id)
    timestamp = quote.created_at.strftime('%Y-%m-%d_%H-%M')
    filename = f"Proposition_commerciale_{quote.opportunity_number}_{timestamp}.{doc_type}"
    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(file_path):
        return jsonify({"error": "Document not found"}), 404

    return send_file(file_path, as_attachment=True)
